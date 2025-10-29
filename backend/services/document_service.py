import os
import re
from datetime import datetime, date
import zipfile
from io import BytesIO
from typing import Dict, Tuple, Any, Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
from typing import List


def _format_date_dmy(dt_str: Optional[str]) -> str:
    """Format a date string to DD/MM/YYYY. Accepts ISO (YYYY-MM-DD) or already formatted inputs."""
    if not dt_str:
        return ''
    try:
        # Try ISO format first
        dt = datetime.fromisoformat(dt_str).date()
        return dt.strftime('%d/%m/%Y')
    except Exception:
        # Try common alternate formats
        for fmt in ('%d/%m/%Y', '%Y/%m/%d', '%d-%m-%Y', '%m/%d/%Y'):
            try:
                dt = datetime.strptime(dt_str, fmt).date()
                return dt.strftime('%d/%m/%Y')
            except Exception:
                continue
        return dt_str


def _set_paragraph_bidi(paragraph, bidi: bool):
    try:
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        # Add or update <w:bidi w:val="true|false"/>
        bidi_elem = pPr.find(qn('w:bidi'))
        if bidi_elem is None:
            bidi_elem = OxmlElement('w:bidi')
            pPr.append(bidi_elem)
        bidi_elem.set(qn('w:val'), 'true' if bidi else 'false')
    except Exception:
        pass


def _replace_in_paragraph(paragraph, replacements: Dict[str, str]):
    """Replace placeholders in a paragraph, robust against split runs.

    Strategy: Concatenate all run texts, perform replacements on the combined
    text, then set the paragraph text in a single run. This may simplify
    formatting for that paragraph but ensures placeholders are replaced even if
    Word split them across runs for styling.
    """
    if not paragraph.runs:
        return
    
    # Store original font size from first run (for footer preservation)
    original_font_size = None
    if paragraph.runs:
        try:
            original_font_size = paragraph.runs[0].font.size
        except:
            pass
    
    combined = ''.join(run.text or '' for run in paragraph.runs)
    replaced = combined
    contains_arabic = False
    for ph, val in replacements.items():
        if ph in replaced:
            val_str = val or ''
            # If value is LTR among Arabic text, add LRM to avoid flipping
            if re.search(r"[\u0600-\u06FF]", replaced):
                val_str = _wrap_ltr_for_arabic_context(val_str)
            replaced = replaced.replace(ph, val_str)
    # Detect Arabic characters in final text (basic range)
    try:
        contains_arabic = re.search(r"[\u0600-\u06FF]", replaced) is not None
    except Exception:
        contains_arabic = False
    if replaced != combined:
        # Safest way: assign paragraph.text which resets runs internally
        try:
            paragraph.text = replaced
            # Restore original font size if it was 8pt (footer) or preserve existing size
            if paragraph.runs and original_font_size:
                for run in paragraph.runs:
                    run.font.size = original_font_size
            # Set bidi direction when Arabic is present
            _set_paragraph_bidi(paragraph, contains_arabic)
            if contains_arabic:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        except Exception:
            # Fallback: set first run and blank the rest
            if paragraph.runs:
                paragraph.runs[0].text = replaced
                for r in paragraph.runs[1:]:
                    r.text = ''
                # Restore font size
                if original_font_size:
                    paragraph.runs[0].font.size = original_font_size
                # Set bidi when Arabic present
                _set_paragraph_bidi(paragraph, contains_arabic)
                if contains_arabic:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _replace_in_table(table, replacements: Dict[str, str]):
    for row in table.rows:
        for cell in row.cells:
            _replace_in_block(cell, replacements)


def _replace_in_block(container, replacements: Dict[str, str]):
    # Paragraphs
    for para in getattr(container, 'paragraphs', []):
        _replace_in_paragraph(para, replacements)
    # Tables
    for table in getattr(container, 'tables', []):
        _replace_in_table(table, replacements)
    # Headers/Footers on sections if present
    for section in getattr(container, 'sections', []):
        try:
            _replace_in_block(section.header, replacements)
        except Exception:
            pass
        try:
            _replace_in_block(section.footer, replacements)
        except Exception:
            pass


def _replace_country_fuzzy(container, country_value: str):
    """Extra safety replacement for (Country) allowing minor spacing/case variations.

    Only applies when a non-empty country_value is provided.
    """
    try:
        if not country_value:
            return
        ws = r"[ \t\r\n\u00A0\u2007\u202F\u2009\u200A\u200B\u2060\uFEFF]*"
        pattern = re.compile(rf"\({ws}C{ws}o{ws}u{ws}n{ws}t{ws}r{ws}y{ws}\)", flags=re.IGNORECASE)

        def _apply(text: str) -> str:
            try:
                return pattern.sub(country_value, text or '')
            except Exception:
                return text

        # Paragraphs
        for para in getattr(container, 'paragraphs', []):
            if para.text and pattern.search(para.text):
                # Replace on combined text; preserve simple formatting by resetting text
                para.text = _apply(para.text)
        # Tables
        for table in getattr(container, 'tables', []):
            for row in table.rows:
                for cell in row.cells:
                    _replace_country_fuzzy(cell, country_value)
        # Sections (headers/footers)
        for section in getattr(container, 'sections', []):
            try:
                _replace_country_fuzzy(section.header, country_value)
            except Exception:
                pass
            try:
                _replace_country_fuzzy(section.footer, country_value)
            except Exception:
                pass
    except Exception:
        pass


def _count_country_placeholders(container) -> int:
    try:
        ws = r"[ \t\r\n\u00A0\u2007\u202F\u2009\u200A\u200B\u2060\uFEFF]*"
        pattern = re.compile(rf"\({ws}C{ws}o{ws}u{ws}n{ws}t{ws}r{ws}y{ws}\)", flags=re.IGNORECASE)
        count = 0
        for para in getattr(container, 'paragraphs', []):
            try:
                if para.text and pattern.search(para.text):
                    count += len(pattern.findall(para.text))
            except Exception:
                continue
        for table in getattr(container, 'tables', []):
            for row in table.rows:
                for cell in row.cells:
                    count += _count_country_placeholders(cell)
        for section in getattr(container, 'sections', []):
            try:
                count += _count_country_placeholders(section.header)
            except Exception:
                pass
            try:
                count += _count_country_placeholders(section.footer)
            except Exception:
                pass
        return count
    except Exception:
        return 0


def _xml_force_replace_country(doc: Document, country_value: str) -> None:
    """Aggressively replace '(Country)' inside XML textboxes and any w:t nodes.

    This handles cases where content resides inside shapes/textboxes (w:txbxContent)
    or structured document tags (w:sdt) that python-docx may not expose via
    document.paragraphs.
    Strategy:
    - For every w:txbxContent//w:p, rebuild the paragraph runs from its concatenated
      text if the placeholder pattern is detected.
    - Additionally, pass through all standalone w:t nodes and apply an in-node replace
      for simple cases.
    Note: This may drop fine-grained formatting inside those paragraphs, by design.
    """
    try:
        if not country_value:
            return
        root = doc._part.element  # lxml element
        ns = root.nsmap.copy() if root.nsmap else {}
        if 'w' not in ns:
            ns['w'] = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        pattern = re.compile(r"\(\s*Country\s*\)", flags=re.IGNORECASE)

        # 1) Replace inside textboxes by rebuilding paragraphs
        for txbx in root.xpath('.//w:txbxContent', namespaces=ns):
            for p in txbx.xpath('.//w:p', namespaces=ns):
                # Collect all texts in order
                t_nodes: List = p.xpath('.//w:r/w:t', namespaces=ns)
                combined = ''.join([t.text or '' for t in t_nodes])
                if not combined:
                    continue
                if not pattern.search(combined):
                    continue
                replaced = pattern.sub(country_value, combined)

                # Remove all existing runs under this paragraph
                for r in p.xpath('./w:r', namespaces=ns):
                    p.remove(r)
                # Build a single run with the replaced text
                r = OxmlElement('w:r')
                t = OxmlElement('w:t')
                # Preserve spaces if value starts/ends with spaces
                t.set(qn('xml:space'), 'preserve')
                t.text = replaced
                r.append(t)
                p.append(r)

        # 2) Global paragraph rebuild for ALL paragraphs (covers w:sdt and any container)
        for p in root.xpath('.//w:p', namespaces=ns):
            t_nodes: List = p.xpath('.//w:r/w:t', namespaces=ns)
            if not t_nodes:
                continue
            combined = ''.join([t.text or '' for t in t_nodes])
            if not combined or not pattern.search(combined):
                continue
            replaced = pattern.sub(country_value, combined)
            # Remove existing runs and rebuild with one run/text
            for r in p.xpath('./w:r', namespaces=ns):
                p.remove(r)
            r = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = replaced
            r.append(t)
            p.append(r)

        # 3) Simple pass: replace inside any individual w:t nodes (covers non-split cases)
        for t in root.xpath('.//w:t', namespaces=ns):
            try:
                if t.text and '(Country)' in t.text:
                    t.text = t.text.replace('(Country)', country_value)
                # Also cover minor spacing variants
                if t.text:
                    t.text = re.sub(rf"\({ws}C{ws}o{ws}u{ws}n{ws}t{ws}r{ws}y{ws}\)", country_value, t.text, flags=re.IGNORECASE)
            except Exception:
                continue
    except Exception:
        # Fail-safe: never raise from aggressive XML step
        pass


def _zip_force_replace_country(docx_path: str, country_value: str) -> None:
    """Open a saved .docx and replace '(Country)' variants across XML parts.

    This is a last-resort approach that ensures replacements even if content lives in
    uncommon containers not traversed by python-docx. It scans these parts:
    - word/document.xml, word/header*.xml, word/footer*.xml
    """
    try:
        print(f"[ZIP_REPLACE] Starting with path='{docx_path}', country='{country_value}'")
        if not (docx_path and country_value and os.path.exists(docx_path)):
            print(f"[ZIP_REPLACE] Validation failed: path={bool(docx_path)}, country={bool(country_value)}, exists={os.path.exists(docx_path) if docx_path else False}")
            return
        ws = r"[ \t\r\n\u00A0\u2007\u202F\u2009\u200A\u200B\u2060\uFEFF]*"
        pattern = re.compile(rf"\({ws}C{ws}o{ws}u{ws}n{ws}t{ws}r{ws}y{ws}\)", flags=re.IGNORECASE)
        print(f"[ZIP_REPLACE] Pattern compiled successfully")

        replacements_made = 0
        with zipfile.ZipFile(docx_path, 'r') as zin:
            out_buf = BytesIO()
            with zipfile.ZipFile(out_buf, 'w', compression=zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename.startswith('word/') and (
                        item.filename == 'word/document.xml' or
                        item.filename.startswith('word/header') or
                        item.filename.startswith('word/footer')
                    ):
                        try:
                            xml_text = data.decode('utf-8')
                            before_count = len(pattern.findall(xml_text))
                            if before_count > 0:
                                print(f"[ZIP_REPLACE] Found {before_count} placeholder(s) in {item.filename}")
                            xml_text = pattern.sub(country_value, xml_text)
                            after_count = len(pattern.findall(xml_text))
                            if before_count > 0:
                                print(f"[ZIP_REPLACE] After replacement: {after_count} placeholder(s) remain in {item.filename}")
                                replacements_made += (before_count - after_count)
                            zout.writestr(item, xml_text.encode('utf-8'))
                        except Exception as e:
                            print(f"[ZIP_REPLACE] Failed to process {item.filename}: {e}")
                            # Keep original if decode fails
                            zout.writestr(item, data)
                            continue
                    else:
                        # Copy unchanged
                        zout.writestr(item, data)
            # Overwrite original file
            with open(docx_path, 'wb') as f:
                f.write(out_buf.getvalue())
        print(f"[ZIP_REPLACE] Completed. Total replacements made: {replacements_made}")
    except Exception as e:
        print(f"[ZIP_REPLACE] Exception occurred: {e}")
        # Never raise; this is a best-effort pass
        pass


def _contains_latin_or_digits(text: str) -> bool:
    if not text:
        return False
    return re.search(r"[A-Za-z0-9]", text) is not None


def _wrap_ltr_for_arabic_context(value: str) -> str:
    """Wrap LTR snippets with LRM marks so they don't break RTL flow.

    Uses Unicode LEFT-TO-RIGHT MARK (\u200E) around latin/digit strings.
    """
    if not value:
        return value
    if _contains_latin_or_digits(value):
        lrm = '\u200E'
        return f"{lrm}{value}{lrm}"
    return value


def _force_container_rtl(container) -> None:
    """Force paragraphs in a container to RTL alignment and bidi.

    Applied for Arabic documents to ensure consistent RTL layout, even for
    paragraphs that end up containing only LTR characters after replacement.
    """
    for para in getattr(container, 'paragraphs', []):
        _set_paragraph_bidi(para, True)
        try:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        except Exception:
            pass
    for table in getattr(container, 'tables', []):
        for row in table.rows:
            for cell in row.cells:
                _force_container_rtl(cell)
    for section in getattr(container, 'sections', []):
        try:
            _force_container_rtl(section.header)
        except Exception:
            pass
        try:
            _force_container_rtl(section.footer)
        except Exception:
            pass


def _extract_first_name(full_name: Optional[str]) -> str:
    """Best-effort extraction of first name from a full name string.

    - Strips common titles (Mr., Mrs., Ms., Dr., Eng, Prof)
    - Splits on whitespace and returns the first non-empty token
    """
    if not full_name:
        return ''
    name = str(full_name).strip()
    if not name:
        return ''
    # Remove common titles at the beginning
    leading = re.sub(r'^\b(mr|mrs|ms|dr|eng|prof|sir|madam|mrs\.|mr\.|ms\.|dr\.|eng\.|prof\.)\b\.?\s+', '', name, flags=re.IGNORECASE)
    # Split by whitespace
    parts = [p for p in re.split(r"\s+", leading) if p]
    if not parts:
        return ''
    return parts[0]


def _resolve_company_country(company: Dict) -> str:
    """Resolve the '(Company Country)' placeholder from company name.

    Mapping rules provided by user:
    - ALOROD AL TAQADAMIAH LEL TASMEM CO -> Jordan
    - Prezlab FZ LLC - Regional Office -> Jordan
    - Prezlab FZ LLC -> Dubai
    - Prezlab Advanced Design Company -> Saudi Arabia
    - Prezlab Digital Design Firm L.L.C. - O.P.C -> Abu Dhabi
    Fallback: company.street (if present) else empty string.
    """
    try:
        name_raw = str(company.get('name') or '').strip()
        name_norm = re.sub(r"\s+", " ", name_raw).lower()

        # Direct mappings (normalized)
        mapping = {
            'alorod al taqadamiah lel tasmem co': 'Jordan',
            'prezlab fz llc - regional office': 'Jordan',
            'prezlab fz llc': 'Dubai',
            'prezlab advanced design company': 'Saudi Arabia',
            'prezlab digital design firm l.l.c. - o.p.c': 'Abu Dhabi',
        }

        # Handle minor variants: if both tokens present in name
        if 'prezlab fz llc' in name_norm and 'regional office' in name_norm:
            return 'Jordan'

        if name_norm in mapping:
            return mapping[name_norm]

        # Fallback to street (previous behavior)
        street = str(company.get('street') or '').strip()
        return street
    except Exception:
        return str(company.get('street') or '')


class DocumentService:
    """Service responsible for generating documents from templates and Odoo data."""

    def __init__(self, odoo_service, employee_service):
        self.odoo_service = odoo_service
        self.employee_service = employee_service

        # Resolve template directory relative to the project root for cross-platform compatibility
        templates_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', 'templates'))

        # Employment letter templates (English)
        self.template_male = os.path.join(templates_root, "Employment Letter - Male.docx")
        self.template_female = os.path.join(templates_root, "Employment Letter - Female.docx")

        # Employment letter templates (Arabic)
        self.template_ar_male = os.path.join(templates_root, "Employment Letter - ARABIC - Male.docx")
        self.template_ar_female = os.path.join(templates_root, "Employment Letter - ARABIC - Female.docx")

        # Generic templates (fallbacks when gender is not available)
        self.template_generic_en = os.path.join(templates_root, "Employment Letter .docx")
        self.template_generic_ar = os.path.join(templates_root, "Employment Letter - ARABIC.docx")

        # Downloads directory inside frontend static folder
        self.downloads_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', 'frontend', 'static', 'downloads', 'employment_letters'))
        
        # Embassy letter templates (English)
        self.template_embassy_male = os.path.join(templates_root, "Employment Letter to Embassies - Male.docx")
        self.template_embassy_female = os.path.join(templates_root, "Employment Letter to Embassies - Female.docx")
        self.template_embassy_generic = os.path.join(templates_root, "Employment Letter to Embassies.docx")
        # Embassy letters downloads directory
        self.embassy_downloads_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', 'frontend', 'static', 'downloads', 'embassy_letters'))

        # Experience letter templates (English)
        self.template_experience_male = os.path.join(templates_root, "Experience Letter - Male.docx")
        self.template_experience_female = os.path.join(templates_root, "Experience Letter - Female.docx")
        self.template_experience_generic = os.path.join(templates_root, "Experience Letter.docx")
        # Experience letters downloads directory
        self.experience_downloads_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', 'frontend', 'static', 'downloads', 'experience_letters'))

        # Service Agreement (Jordan) template
        self.template_service_agreement = os.path.join(templates_root, "Prezlab-Jordan- Service.docx")
        # Service agreements downloads directory
        self.service_downloads_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', 'frontend', 'static', 'downloads', 'service_agreements'))

        # Company-specific service agreement templates
        # Defaults for Jordan-based companies remain on the original Jordan template
        self.template_service_agreement_map = {
            'Prezlab FZ LLC - Regional Office': self.template_service_agreement,
            'ALOROD AL TAQADAMIAH LEL TASMEM CO': self.template_service_agreement,
            'Prezlab Advanced Design Company': os.path.join(templates_root, "KSA Service Agreement - (Prezlab Advanced Design Company).docx"),
            'Prezlab FZ LLC': os.path.join(templates_root, "Dubai Service Agreement - (Prezlab FZ LLC).docx"),
            'Prezlab Digital Design Firm L.L.C. - O.P.C': os.path.join(templates_root, "Abu Dhabi Service Agreement - (Prezlab Digital Design Firm).docx"),
        }

    def _ensure_downloads_dir(self):
        os.makedirs(self.downloads_dir, exist_ok=True)
        os.makedirs(self.embassy_downloads_dir, exist_ok=True)
        os.makedirs(self.experience_downloads_dir, exist_ok=True)
        os.makedirs(self.service_downloads_dir, exist_ok=True)

    def _read_employee_with_fields(self, employee_id: int, fields: list) -> Tuple[bool, Any]:
        # Use safe employee read to gracefully handle field-level AccessError (drops forbidden fields)
        return self.employee_service._safe_employee_read([employee_id], fields)

    def _read_company_with_fields(self, company_id: int, fields: list) -> Tuple[bool, Any]:
        params = {
            'args': [[company_id]],
            'kwargs': {'fields': fields}
        }
        return self.employee_service._make_odoo_request('res.company', 'read', params)

    def _build_replacements(self, employee: Dict, company: Dict, lang: str = 'en') -> Dict[str, str]:
        current_date = date.today().strftime('%d/%m/%Y')

        joining_date_raw = employee.get('x_studio_joining_date')
        joining_date = _format_date_dmy(joining_date_raw) if joining_date_raw else ''

        first_last_name = employee.get('name') or ''
        first_name_only = _extract_first_name(first_last_name)
        position = employee.get('job_title') or ''
        work_location_val = employee.get('x_studio_work_location_country')
        # Handle many2one/list values like [id, 'Name']
        if isinstance(work_location_val, (list, tuple)) and len(work_location_val) > 1:
            work_address_en = str(work_location_val[1])
        else:
            work_address_en = str(work_location_val or '')
        arabic_work_address = work_address_en  # default to English as requested

        company_name_en = company.get('name') or ''
        company_name_ar = company.get('arabic_name') or company_name_en
        # Resolve Company Country using explicit mapping first, fallback to street
        company_street = _resolve_company_country(company)
        company_registry = str(company.get('company_registry') or '')

        # Department name (many2one)
        dept_val = employee.get('department_id')
        if isinstance(dept_val, (list, tuple)) and len(dept_val) > 1:
            department_name = str(dept_val[1])
        else:
            department_name = str(dept_val or '')

        replacements = {
            '(Current Date)': current_date,
            '(First and Last Name)': first_last_name,
            '(First Name)': first_name_only,
            '(Department)': department_name,
            '(Company)': company_name_en,
            '(Company Country)': company_street,
            '(CR)': company_registry,
            '(DD/MM/YYYY)': joining_date,
            '(Position)': position,
            '(P&C)': 'Faisal Abdullah AlMamun',
            '(Work address)': work_address_en,
            '(Arabic Work address)': arabic_work_address,
            '(CompanyA)': company_name_ar,
        }
        if (lang or 'en').lower().startswith('ar'):
            arabic_full_name = employee.get('x_studio_employee_arabic_name') or ''
            # Arabic template specific placeholder
            replacements['(الاسم الكامل)'] = arabic_full_name
        return replacements

    def generate_employment_letter(self, lang: str = 'en') -> Tuple[bool, Any]:
        """Generate an Employment Letter for the current user and return metadata including file URL."""
        # Ensure active Odoo session
        session_ok, session_msg = self.odoo_service.ensure_active_session()
        if not session_ok:
            return False, f'Odoo session error: {session_msg}'

        # Get current user's employee data to find employee id
        emp_ok, emp_data = self.employee_service.get_current_user_employee_data()
        if not emp_ok or not emp_data:
            return False, f'Could not fetch employee data: {emp_data}'

        employee_id = emp_data.get('id')
        if not employee_id:
            return False, 'Employee ID not found for current user'

        # Read required employee fields (ensures custom fields are fetched reliably)
        emp_fields = ['name', 'gender', 'x_studio_rf_gender', 'job_title', 'department_id', 'company_id', 'x_studio_joining_date', 'x_studio_work_location_country', 'x_studio_employee_arabic_name']
        ok, emp_read = self._read_employee_with_fields(employee_id, emp_fields)
        if not ok or not emp_read:
            return False, f'Failed to read employee fields: {emp_read}'
        employee = emp_read[0] if isinstance(emp_read, list) else emp_read

        # Resolve company id
        company_id = None
        comp_val = employee.get('company_id')
        if isinstance(comp_val, list) and comp_val:
            company_id = comp_val[0]
        elif isinstance(comp_val, int):
            company_id = comp_val

        if not company_id:
            return False, 'Company not found for employee'

        # Read required company fields
        comp_fields = ['name', 'street', 'company_registry', 'arabic_name']
        ok, comp_read = self._read_company_with_fields(company_id, comp_fields)
        if not ok or not comp_read:
            return False, f'Failed to read company fields: {comp_read}'
        company = comp_read[0] if isinstance(comp_read, list) else comp_read

        # Choose template based on gender and language; use generic fallback when gender missing/unknown
        gender_raw = (employee.get('gender') or '').strip()
        if not gender_raw:
            gender_raw = str(employee.get('x_studio_rf_gender') or '').strip()
        gender = gender_raw.lower()
        is_female = gender.startswith('f')
        is_male = gender.startswith('m')
        if (lang or 'en').lower().startswith('ar'):
            if is_female:
                template_path = self.template_ar_female
            elif is_male:
                template_path = self.template_ar_male
            else:
                template_path = self.template_generic_ar
        else:
            if is_female:
                template_path = self.template_female
            elif is_male:
                template_path = self.template_male
            else:
                template_path = self.template_generic_en
        if not os.path.exists(template_path):
            return False, f'Template not found: {template_path}'

        # Prepare replacements
        replacements = self._build_replacements(employee, company, lang=lang)

        # Load, replace, and save document
        doc = Document(template_path)
        # For Arabic docs, force container RTL before and after replacements
        if (lang or 'en').lower().startswith('ar'):
            _force_container_rtl(doc)
        _replace_in_block(doc, replacements)
        if (lang or 'en').lower().startswith('ar'):
            _force_container_rtl(doc)

        # Ensure downloads dir exists
        self._ensure_downloads_dir()

        # Build filename in the format: "[doc type] - [person name].docx"
        # Example: "Arabic Employment Letter - Omar basem elhasan.docx"
        def _sanitize_filename_part(text: str) -> str:
            # Remove characters not allowed in filenames on Windows and most filesystems
            return re.sub(r'[\\\\/:*?"<>|]', '-', text).strip()

        person_name_display = employee.get('name') or 'Employee'
        is_ar = (lang or 'en').lower().startswith('ar')
        doc_type_display = 'arabic employment letter' if is_ar else 'employment letter'
        filename = f"{_sanitize_filename_part(doc_type_display)} - {_sanitize_filename_part(person_name_display)}.docx"
        filepath = os.path.join(self.downloads_dir, filename)
        doc.save(filepath)

        # Build URL relative to Flask static
        file_url = f"/static/downloads/employment_letters/{filename}"

        return True, {
            'file_name': filename,
            'file_url': file_url,
            'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        }

    def generate_service_agreement(self, full_name: str, private_street: str = '', company_name: str = '') -> Tuple[bool, Any]:
        """Generate a Service Agreement using the provided template and return attachment metadata.

        Placeholders replaced:
        - (First and Last Name)
        - (Current Date)
        """
        try:
            # Resolve template by company when available
            # Note: Backward compatibility – if caller doesn't pass company_name, default mapping uses Jordan template
            company_name = locals().get('company_name', '')  # will be injected via updated signature below
            template_path = self.template_service_agreement_map.get(str(company_name).strip(), self.template_service_agreement)
            if not os.path.exists(template_path):
                return False, f"Template not found: {template_path}"

            self._ensure_downloads_dir()

            today_str = date.today().strftime('%d/%m/%Y')
            replacements = {
                '(First and Last Name)': str(full_name or '').strip(),
                '(Current Date)': today_str,
            }

            doc = Document(template_path)
            # Include optional Private Street placeholder if provided
            if private_street:
                replacements['(Private Street)'] = str(private_street)
            _replace_in_block(doc, replacements)

            # Filename: Service_Agreement_[Name]_[YYYYMMDD].docx
            def _sanitize_filename_part(text: str) -> str:
                import re as _re
                return _re.sub(r'[\\/:*?"<>|]', '-', str(text or '')).strip()

            filename = f"Service_Agreement_{_sanitize_filename_part(full_name)}_{date.today().strftime('%Y%m%d')}.docx"
            filepath = os.path.join(self.service_downloads_dir, filename)
            doc.save(filepath)

            file_url = f"/static/downloads/service_agreements/{filename}"
            return True, {
                'file_name': filename,
                'file_url': file_url,
                'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            }
        except Exception as e:
            return False, f"Service agreement generation error: {e}"

    def generate_experience_letter(self) -> Tuple[bool, Any]:
        """Generate an Experience Letter for the current user and return metadata including file URL."""
        # Ensure active Odoo session
        session_ok, session_msg = self.odoo_service.ensure_active_session()
        if not session_ok:
            return False, f'Odoo session error: {session_msg}'

        # Get current user's employee data to find employee id
        emp_ok, emp_data = self.employee_service.get_current_user_employee_data()
        if not emp_ok or not emp_data:
            return False, f'Could not fetch employee data: {emp_data}'

        employee_id = emp_data.get('id')
        if not employee_id:
            return False, 'Employee ID not found for current user'

        # Read required employee fields (extend if templates require more)
        emp_fields = ['name', 'gender', 'x_studio_rf_gender', 'job_title', 'department_id', 'company_id', 'x_studio_joining_date', 'x_studio_work_location_country', 'x_studio_employee_arabic_name']
        ok, emp_read = self._read_employee_with_fields(employee_id, emp_fields)
        if not ok or not emp_read:
            return False, f'Failed to read employee fields: {emp_read}'
        employee = emp_read[0] if isinstance(emp_read, list) else emp_read

        # Resolve company id
        company_id = None
        comp_val = employee.get('company_id')
        if isinstance(comp_val, list) and comp_val:
            company_id = comp_val[0]
        elif isinstance(comp_val, int):
            company_id = comp_val
        if not company_id:
            return False, 'Company not found for employee'

        # Read required company fields
        comp_fields = ['name', 'street', 'company_registry', 'arabic_name']
        ok, comp_read = self._read_company_with_fields(company_id, comp_fields)
        if not ok or not comp_read:
            return False, f'Failed to read company fields: {comp_read}'
        company = comp_read[0] if isinstance(comp_read, list) else comp_read

        # Choose template by gender with generic fallback when gender missing/unknown
        gender_raw = (employee.get('gender') or '').strip()
        if not gender_raw:
            gender_raw = str(employee.get('x_studio_rf_gender') or '').strip()
        gender = gender_raw.lower()
        is_female = gender.startswith('f')
        is_male = gender.startswith('m')
        if is_female:
            template_path = self.template_experience_female
        elif is_male:
            template_path = self.template_experience_male
        else:
            template_path = self.template_experience_generic
        if not os.path.exists(template_path):
            return False, f'Template not found: {template_path}'

        # Prepare replacements using base mapping
        replacements = self._build_replacements(employee, company, lang='en')

        # Build document
        doc = Document(template_path)
        _replace_in_block(doc, replacements)

        # Ensure directories exist
        self._ensure_downloads_dir()

        # Filename: "experience letter - Name.docx"
        def _sanitize_filename_part(text: str) -> str:
            return re.sub(r'[\\/:*?"<>|]', '-', text).strip()

        person_name_display = employee.get('name') or 'Employee'
        filename = f"experience letter - {_sanitize_filename_part(person_name_display)}.docx"
        filepath = os.path.join(self.experience_downloads_dir, filename)
        doc.save(filepath)

        file_url = f"/static/downloads/experience_letters/{filename}"
        return True, {
            'file_name': filename,
            'file_url': file_url,
            'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        }

    def generate_embassy_letter(self, country: str, start_date: str, end_date: str) -> Tuple[bool, Any]:
        """Generate an Employment Letter to Embassies with country and date range placeholders."""
        print(f"[EMBASSY] generate_embassy_letter called with country='{country}', start_date='{start_date}', end_date='{end_date}'")
        # Ensure active Odoo session
        session_ok, session_msg = self.odoo_service.ensure_active_session()
        if not session_ok:
            return False, f'Odoo session error: {session_msg}'

        # Fetch current user employee data
        emp_ok, emp_data = self.employee_service.get_current_user_employee_data()
        if not emp_ok or not emp_data:
            return False, f'Could not fetch employee data: {emp_data}'

        employee_id = emp_data.get('id')
        if not employee_id:
            return False, 'Employee ID not found for current user'

        # Read required employee fields
        emp_fields = ['name', 'gender', 'x_studio_rf_gender', 'job_title', 'department_id', 'company_id', 'x_studio_joining_date', 'x_studio_work_location_country', 'x_studio_employee_arabic_name']
        ok, emp_read = self._read_employee_with_fields(employee_id, emp_fields)
        if not ok or not emp_read:
            return False, f'Failed to read employee fields: {emp_read}'
        employee = emp_read[0] if isinstance(emp_read, list) else emp_read

        # Resolve company
        company_id = None
        comp_val = employee.get('company_id')
        if isinstance(comp_val, list) and comp_val:
            company_id = comp_val[0]
        elif isinstance(comp_val, int):
            company_id = comp_val
        if not company_id:
            return False, 'Company not found for employee'

        comp_fields = ['name', 'street', 'company_registry', 'arabic_name']
        ok, comp_read = self._read_company_with_fields(company_id, comp_fields)
        if not ok or not comp_read:
            return False, f'Failed to read company fields: {comp_read}'
        company = comp_read[0] if isinstance(comp_read, list) else comp_read

        # Choose template by gender with generic fallback when gender missing/unknown
        gender_raw = (employee.get('gender') or '').strip()
        if not gender_raw:
            gender_raw = str(employee.get('x_studio_rf_gender') or '').strip()
        gender = gender_raw.lower()
        is_female = gender.startswith('f')
        is_male = gender.startswith('m')
        if is_female:
            template_path = self.template_embassy_female
        elif is_male:
            template_path = self.template_embassy_male
        else:
            template_path = self.template_embassy_generic
        if not os.path.exists(template_path):
            return False, f'Template not found: {template_path}'

        # Prepare replacements (start with base then add embassy placeholders)
        replacements = self._build_replacements(employee, company, lang='en')
        replacements.update({
            '(Start Date)': _format_date_dmy(start_date),
            '(End Date)': _format_date_dmy(end_date)
        })
        # Only replace (Country) if a country was explicitly provided; otherwise leave the placeholder in the template
        country_display = (country or '').strip()
        print(f"[EMBASSY] country_display after strip: '{country_display}'")
        if country_display:
            replacements['(Country)'] = country_display
            print(f"[EMBASSY] Added '(Country)' -> '{country_display}' to replacements")
        else:
            print(f"[EMBASSY] WARNING: country_display is empty, not replacing (Country)")

        # Build document
        doc = Document(template_path)
        print(f"[EMBASSY] Calling _replace_in_block with {len(replacements)} replacements")
        _replace_in_block(doc, replacements)

        # Ensure directories exist
        self._ensure_downloads_dir()

        # Filename: "embassy employment letter - Name.docx"
        def _sanitize_filename_part(text: str) -> str:
            return re.sub(r'[\\/:*?"<>|]', '-', text).strip()

        person_name_display = employee.get('name') or 'Employee'
        filename = f"embassy employment letter - {_sanitize_filename_part(person_name_display)}.docx"
        filepath = os.path.join(self.embassy_downloads_dir, filename)
        print(f"[EMBASSY] Saving document to: {filepath}")
        doc.save(filepath)
        # Final safety: zip-level pass over XML parts
        if country_display:
            print(f"[EMBASSY] Calling _zip_force_replace_country with country='{country_display}'")
            _zip_force_replace_country(filepath, country_display)
            print(f"[EMBASSY] _zip_force_replace_country completed")
        else:
            print(f"[EMBASSY] Skipping zip replacement because country_display is empty")

        file_url = f"/static/downloads/embassy_letters/{filename}"
        return True, {
            'file_name': filename,
            'file_url': file_url,
            'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        }


